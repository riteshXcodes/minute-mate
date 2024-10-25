import time
from flask import Flask, redirect, render_template, request, jsonify
from flask_socketio import SocketIO, join_room, leave_room
import uuid
from docx import Document
from dotenv import load_dotenv
# from datetime import datetime
import os
from langchain import LLMChain, PromptTemplate
from langchain.llms import DeepInfra
from datetime import datetime


app = Flask(__name__)
socketio = SocketIO(app, 
                    cors_allowed_origins="*",
                    #max_http_buffer_size=9999999999
                    )

rooms = {}  # For simplicity, store rooms in-memory

load_dotenv()

DEEPINFRA_API_TOKEN = os.getenv('DEEPINFRA_API_TOKEN')

def save_summary_and_actions_to_doc(summary_text, action_items,new_doc_path):
    """Save the summary text and action items to a .docx file."""
    doc = Document()
    doc.add_heading('Meeting Summary', 0)
    doc.add_paragraph(summary_text)

    doc.add_heading('Action Items', level=1)
    for idx, item in enumerate(action_items, 1):
        doc.add_paragraph(f"{idx}. {item}")

    # doc_path = 'meeting_summary_with_action_items.docx'
    doc_path = new_doc_path
    doc.save(doc_path)
    return doc_path


    
@app.route('/summarize', methods=['POST'])
def generate_summary_and_action_items():
    try:
        # Get the meeting transcript from the POST request
        text_chunk = request.json['transcript']

        # Define the template for generating the summary
        summary_template = """
        Write a concise summary of the text, return your responses with 5 lines that cover the key points of the text.
        ```{text}```
        SUMMARY:
        """
        
        # Create the prompt for summary
        summary_prompt = PromptTemplate(template=summary_template, input_variables=["text"])
        summary_chain = LLMChain(prompt=summary_prompt, llm=DeepInfra(model_id="meta-llama/Llama-2-70b-chat-hf"))

        # Generate the summary
        summary = summary_chain.invoke(text_chunk)
        if not summary:
            raise Exception("Summary generation failed")

        # Define the template for generating action items
        action_items_template = """
        Generate a list of actionable items from the following text, return your responses with bullet points.
        ```{text}```
        ACTION ITEMS:
        """
        
        # Create the prompt for action items
        action_items_prompt = PromptTemplate(template=action_items_template, input_variables=["text"])
        action_items_chain = LLMChain(prompt=action_items_prompt, llm=DeepInfra(model_id="meta-llama/Llama-2-70b-chat-hf"))

        # Generate the action items
        action_items_text = action_items_chain.invoke(text_chunk)
        if not action_items_text:
            raise Exception("Action items extraction failed")

        # Split the action items correctly
        action_items = [item.strip() for item in action_items_text.split("\n") if item.strip()]

        
        randINT = uuid.uuid4().hex[:8]
        doc_name = f"summary_action_items_{randINT}.docx"
        newPath = os.path.join('C:/Users/Dell/Desktop/repo', doc_name)

        # Save the summary and action items to a .docx file
        doc_path = save_summary_and_actions_to_doc(summary, action_items,newPath)

        return jsonify(summary=summary, action_items=action_items, document_path=doc_path)
    except Exception as e:
        return jsonify(error=str(e)), 500


@app.route('/create_room', methods=['POST'])
def create_room():
    room_id = uuid.uuid4().hex[:8] # short unique id for room
    if room_id in rooms:
        return jsonify(success=False, message="Room already exists"), 400
    rooms[room_id] = {"participants": []}
    return jsonify({
        "success": True,
        "room_id": room_id
    })

@app.route('/get_rooms', methods=['GET'])
def get_rooms():
    return jsonify(list(rooms.keys()))

@app.route('/delete_room/<room_id>', methods=['GET'])
def delete_room(room_id):
    del rooms[room_id]
    return jsonify({
        "success": True,
        "message": "Room deleted"
    })

@socketio.on('join')
def on_join(data):
    username = data['username']
    room_id = data['room_id']
    if room_id not in rooms:
        return {"success": False, "message": "Room not found"}
    join_room(room_id)
    rooms[room_id]["participants"].append(username)

    # Notify other users in the room about the new user
    for participant in rooms[room_id]["participants"]:
        if participant != username:
            socketio.emit('user_joined', {"username": username, "userId": request.sid}, room=participant)
            time.sleep(5)

    print(f"User {username} has joined room {room_id}")

    # Notify other users in the room about the new user
    socketio.emit('user_joined', {"username": username, "userId": request.sid}, room=room_id)
    

@socketio.on('leave')
def on_leave(data):
    username = data['username']
    room_id = data['room_id']
    if room_id not in rooms:
        return {"success": False, "message": "Room not found"}
    leave_room(room_id)
    rooms[room_id]["participants"].remove(username)

    socketio.emit('user_left', {"username": username, "userId": request.sid}, room=room_id)

@socketio.on('signal')
def on_signal(data):
    target_user = data['userId']
    room_id = data['room_id']
    signal = data['signal']

    print(f"Signal from {request.sid} to {target_user}")
    # Send the signal to the specified user in the room
    socketio.emit('signal', {"userId": request.sid, "signal": signal}, room=target_user)


@socketio.on('share_screen')
def handle_share_screen(data):
    room_id = data['room_id']
    user_id = data['userId']
    is_sharing = data['isScreenSharing']
    
    # Broadcast this event to all other users in the same room
    socketio.emit('screen_sharing_status', {'userId': user_id, 'isScreenSharing': is_sharing}, room=room_id, include_self=False)
    
@socketio.on('stop_share_screen')
def handle_stop_share_screen(data):
    room_id = data['room_id']
    user_id = data['userId']
    # Broadcast to all users in the room that the stream has been updated
    socketio.emit('stream_updated', {'userId': user_id, 'type': 'webcam'}, room=room_id)

@socketio.on('request_new_stream')
def on_request_new_stream(data):
    target_user = data['userId']
    room_id = data['room_id']
    
    # Inform the target user to create a new offer for their stream
    socketio.emit('create_new_offer', {'fromUserId': request.sid}, room=target_user)

@socketio.on('transcript_message')
def handle_transcript_message(data):
    username = data['username']
    transcript = data['transcript']
    room_id = data['room_id']
    socketio.emit('broadcast_transcript', {'username': username, 'transcript': transcript}, room=room_id)


@socketio.on('new_message')
def handle_new_message(data):
    message = data['message']
    username = data['username']
    room_id = data['room_id']  # You need to have the concept of rooms

    # Broadcasting the message to all users in the room
    socketio.emit('message_received', {'message': message, 'username': username}, room=room_id, include_self=False)

@app.route('/room_join/<room_id>')
def room_join(room_id):
    if room_id in rooms.keys():
        return render_template('second_page.html', room_id=room_id)
    
    return redirect('/')

@app.route('/')
def index():
    return render_template('first_page.html')

if __name__ == '__main__':
    socketio.run(app, host='127.0.0.1', port=5001, debug=True)
